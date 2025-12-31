
import tkinter as tk
from tkinter import ttk, filedialog
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
from builtins import PendingDeprecationWarning
from PIL import Image, ImageTk
import json

# --- Chart Generators ---
def generate_budget_chart(data):
    labels = list(data.keys())
    values = list(data.values())
    plt.figure(figsize=(5,5))
    plt.pie(values, labels=labels, autopct='%1.1f%%')
    plt.title("Budget Breakdown")
    plt.savefig("budget_chart.png")
    plt.close()

def generate_timeline_chart(timeline):
    milestones = list(timeline.keys())
    months = list(timeline.values())
    plt.figure(figsize=(6,4))
    plt.bar(milestones, months, color='skyblue')
    plt.title("Project Timeline (Months)")
    plt.xlabel("Milestones")
    plt.ylabel("Duration")
    plt.savefig("timeline_chart.png")
    plt.close()

# --- Export Functions ---
def export_pdf(company, problem, solution, budget, timeline):
    generate_budget_chart(budget)
    generate_timeline_chart(timeline)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Business Proposal for {company}", ln=True, align="C")
    pdf.multi_cell(0, 10, f"Problem: {problem}\n\nSolution: {solution}\n")
    pdf.image("budget_chart.png", x=30, y=80, w=150)
    pdf.image("timeline_chart.png", x=30, y=160, w=150)
    pdf.output("proposal.pdf")
    

def export_word(company, problem, solution, budget, timeline):
    doc = Document()
    doc.add_heading(f"Business Proposal for {company}", 0)
    doc.add_paragraph(f"Problem: {problem}")
    doc.add_paragraph(f"Solution: {solution}")
    doc.add_heading("Budget Breakdown", level=1)
    for k,v in budget.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Timeline", level=1)
    for k,v in timeline.items():
        doc.add_paragraph(f"{k}: {v} months")
    doc.save("proposal.docx")
    

# --- Save/Load Draft ---
def save_draft():
    company = company_entry.get()
    problem = problem_entry.get("1.0", tk.END).strip()
    solution = solution_entry.get("1.0", tk.END).strip()
    budget = {
        "Marketing": marketing_entry.get(),
        "Development": dev_entry.get(),
        "Operations": ops_entry.get(),
        "Misc": misc_entry.get()
    }
    timeline = {
        "Planning": plan_entry.get(),
        "Execution": exec_entry.get(),
        "Launch": Launch_entry.get()
    }
    draft = {
        "company": company,
        "problem": problem,
        "solution": solution,
        "budget": budget,
        "timeline": timeline
    }
    file_path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON files","*.json")])
    if file_path:
        with open(file_path, "w") as f:
            json.dump(draft, f)
    status_Label.config(text=f"Draft saved to {file_path}")

def load_draft():
    file_path = filedialog.askopenfilename(filetypes=[("JSON files","*.json")])
    if file_path:
        with open(file_path, "r") as f:
            draft = json.load(f)
        company_entry.delete(0, tk.END)
        company_entry.insert(0, draft["company"])
        problem_entry.delete("1.0", tk.END)
        problem_entry.insert("1.0", draft["problem"])
        solution_entry.delete("1.0", tk.END)
        solution_entry.insert("1.0", draft["solution"])
        marketing_entry.delete(0, tk.END); marketing_entry.insert(0, draft["budget"]["Marketing"])
        dev_entry.delete(0, tk.END); dev_entry.insert(0, draft["budget"]["Development"])
        ops_entry.delete(0, tk.END); ops_entry.insert(0, draft["budget"]["Operations"])
        misc_entry.delete(0, tk.END); misc_entry.insert(0, draft["budget"]["Misc"])
        plan_entry.delete(0, tk.END); plan_entry.insert(0, draft["timeline"]["Planning"])
        exec_entry.delete(0, tk.END); exec_entry.insert(0, draft["timeline"]["Execution"])
        launch_entry.delete(0, tk.END); launch_entry.insert(0, draft["timeline"]["Launch"])
        status_label.config(text=f"Draft loaded from {file_path}")

# --- GUI Logic ---
def build_proposal():
    try:
        company = company_entry.get()
        problem = problem_entry.get("1.0", tk.END).strip()
        solution = solution_entry.get("1.0", tk.END).strip()

        budget = {
            "Marketing": int(marketing_entry.get()),
            "Development": int(dev_entry.get()),
            "Operations": int(ops_entry.get()),
            "Misc": int(misc_entry.get())
        }

        timeline = {
            "Planning": int(plan_entry.get()),
            "Execution": int(exec_entry.get()),
            "Launch": int(launch_entry.get())
        }

        export_pdf(company, problem, solution, budget, timeline)
        export_word(company, problem, solution, budget, timeline)
        status_label.config(text="Proposal exported as proposal.pdf and proposal.docx")
        update_preview(company, problem, solution, budget, timeline)

    except ValueError:
        status_label.config(text="Please fill in all fields with valid numbers.")


    try:
        company = company_entry.get()
        problem = problem_entry.get("1.0", tk.END).strip()
        solution = solution_entry.get("1.0", tk.END).strip()

        # Validate budget inputs
        budget = {
            "Marketing": int(marketing_entry.get()),
            "Development": int(dev_entry.get()),
            "Operations": int(ops_entry.get()),
            "Misc": int(misc_entry.get())
        }

        # Validate timeline inputs
        timeline = {
            "Planning": int(plan_entry.get()),
            "Execution": int(exec_entry.get()),
            "Launch": int(launch_entry.get())
        }

        export_pdf(company, problem, solution, budget, timeline)
        export_word(company, problem, solution, budget, timeline)
        status_label.config(text="Proposal exported as proposal.pdf and proposal.docx")
        update_preview(company, problem, solution, budget, timeline)

    except ValueError:
       status_label.config(text="Please fill in all fields with valid numbers.")


def update_preview(company, problem, solution, budget, timeline):
    preview_text.delete("1.0", tk.END)
    preview_text.insert(tk.END, f"Business Proposal for {company}\n\n")
    preview_text.insert(tk.END, f"Problem:\n{problem}\n\n")
    preview_text.insert(tk.END, f"Solution:\n{solution}\n\n")
    preview_text.insert(tk.END, "Budget Breakdown:\n")
    for k,v in budget.items():
        preview_text.insert(tk.END, f" - {k}: {v}\n")
    preview_text.insert(tk.END, "\nTimeline:\n")
    for k,v in timeline.items():
        preview_text.insert(tk.END, f" - {k}: {v} months\n")

    generate_budget_chart(budget)
    generate_timeline_chart(timeline)

    budget_img = Image.open("budget_chart.png").resize((250,250))
    budget_photo = ImageTk.PhotoImage(budget_img)
    budget_label.config(image=budget_photo)
    budget_label.image = budget_photo

    timeline_img = Image.open("timeline_chart.png").resize((300,200))
    timeline_photo = ImageTk.PhotoImage(timeline_img)
    timeline_label.config(image=timeline_photo)
    timeline_label.image = timeline_photo
    
root = tk.Tk()
tab_control = ttk.Notebook(root)
info_tab = ttk.Frame(tab_control)
budget_tab = ttk.Frame(tab_control)
timeline_tab = ttk.Frame(tab_control)
export_tab = ttk.Frame(tab_control)
preview_tab = ttk.Frame(tab_control)

tab_control.add(info_tab, text="Info")
tab_control.add(budget_tab, text="Budget")
tab_control.add(timeline_tab, text="Timeline")
tab_control.add(export_tab, text="Export")
tab_control.add(preview_tab, text="Preview")
tab_control.pack(expand=1, fill="both")

# Info Tab
ttk.Label(info_tab, text="Company Name").pack()
company_entry = ttk.Entry(info_tab); company_entry.pack()

ttk.Label(info_tab, text="Problem Statement").pack()
problem_entry = tk.Text(info_tab, height=6, width=60, wrap="word"); problem_entry.pack()
scrollbar_problem = ttk.Scrollbar(info_tab, orient="vertical", command=problem_entry.yview)
problem_entry.configure(yscrollcommand=scrollbar_problem.set)
scrollbar_problem.pack(side="right", fill="y")

ttk.Label(info_tab, text="Solution").pack()
solution_entry = tk.Text(info_tab, height=6, width=60, wrap="word"); solution_entry.pack()
scrollbar_solution = ttk.Scrollbar(info_tab, orient="vertical", command=solution_entry.yview)
solution_entry.configure(yscrollcommand=scrollbar_solution.set)
scrollbar_solution.pack(side="right", fill="y")

ttk.Button(export_tab, text="Generate Proposal", command=build_proposal).pack(pady=10)
status_label = ttk.Label(export_tab, text="")
status_label.pack()

preview_text = tk.Text(preview_tab, height=15, width=70, wrap="word")
preview_text.pack()

budget_label = ttk.Label(preview_tab)
budget_label.pack()

timeline_label = ttk.Label(preview_tab)
timeline_label.pack()

    

# Budget Tab
ttk.Label(budget_tab, text="Marketing").pack()
marketing_entry = ttk.Entry(budget_tab); marketing_entry.insert(0,"40"); marketing_entry.pack()
ttk.Label(budget_tab, text="Development").pack()
dev_entry = ttk.Entry(budget_tab); dev_entry.insert(0,"30"); dev_entry.pack()
ttk.Label(budget_tab, text="Operations").pack()
ops_entry = ttk.Entry(budget_tab); ops_entry.insert(0,"20"); ops_entry.pack()
ttk.Label(budget_tab, text="Misc").pack()
misc_entry = ttk.Entry(budget_tab); misc_entry.insert(0,"10"); misc_entry.pack()

# Timeline Tab
ttk.Label(timeline_tab, text="Planning (months)").pack(pady=5)
plan_entry = ttk.Entry(timeline_tab)
plan_entry.insert(0, "2")
plan_entry.pack()

ttk.Label(timeline_tab, text="Execution (months)").pack(pady=5)
exec_entry = ttk.Entry(timeline_tab)
exec_entry.insert(0, "6")  # Add a default value to prevent empty input
exec_entry.pack()

ttk.Label(timeline_tab, text="Launch (months)").pack(pady=5)
launch_entry = ttk.Entry(timeline_tab)
launch_entry.insert(0, "1")
launch_entry.pack()


# --- GUI Setup ---
root = tk.Tk()
root.title("Business Proposal Builder")
root.mainloop()
